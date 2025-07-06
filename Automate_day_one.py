import os
import re
import shutil
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from PyPDF2 import PdfReader
from docx import Document

# ------------------------
# EXCEL EXTRACTION SECTION
# ------------------------

def extract_communication_goals(text, student_name=None):
    comm_blocks = re.findall(
        r"Domain\(s\)/TSAA\(s\):\s*Communication(.*?)(?=\nDomain\(s\)|\nAssessments|\nAccommodations|Page \d|\Z)",
        text, re.DOTALL | re.IGNORECASE
    )
    results = []

    for section in comm_blocks:
        goal_blocks = re.split(r"\bGoal:\s*", section, flags=re.IGNORECASE)
        for block in goal_blocks[1:]:
            goal_text = re.split(
                r"(?:\nShort-term Objectives|Assessment Procedures|Progress Reported|\n[A-Z][a-z]+:|\Z)",
                block, maxsplit=1, flags=re.IGNORECASE
            )[0].strip().replace('\n', ' ')

            benchmark_blocks = re.findall(
                r"Short-term Objectives or Benchmarks:(.*?)(?=\n(?:Short-term Objectives or Benchmarks:|Goal:|Domain\(s\)|Assessments|Accommodations|Page \d|\Z))",
                block, re.DOTALL | re.IGNORECASE
            )

            subgoals = []
            for bblock in benchmark_blocks:
                lines = bblock.strip().splitlines()
                for line in lines:
                    line = line.strip()
                    if (
                        len(line.split()) > 5 and
                        " will " in line.lower() and
                        not re.search(r"(student be|participate|assessment|instruction|comment)", line, re.IGNORECASE)
                    ):
                        subgoals.append(line)

            if not subgoals:
                lines = block.strip().splitlines()
                for line in lines:
                    line = line.strip()
                    if (
                        len(line.split()) > 5 and
                        " will " in line.lower() and
                        not re.search(r"(student be|participate|assessment|instruction|comment)", line, re.IGNORECASE)
                    ):
                        subgoals.append(line)

            if goal_text:
                results.append({
                    "goal": goal_text,
                    "subgoals": subgoals
                })

    return results

def extract_name(text):
    match = re.search(r"Student:\s+([A-Z\s'-]+),\s*([A-Z][a-zA-Z'-]+)", text)
    if match:
        last = match.group(1).strip().title()
        first = match.group(2).strip().title()
        return first, last

    fallback = re.search(r"\b([A-Z][a-z]+)\s+will\b", text)
    if fallback:
        return fallback.group(1), "Student"

    return "Unknown", "Student"

def analyze_pdfs_and_generate_excel(pdf_folder, output_excel_path):
    all_data = []
    max_goals = 0
    max_subgoals_per_goal = {}

    for filename in os.listdir(pdf_folder):
        if filename.lower().endswith(".pdf"):
            path = os.path.join(pdf_folder, filename)
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

            student_id_match = re.search(r"\b(\d{10})\b", text)
            student_id = student_id_match.group(1) if student_id_match else "NoID"

            first_name, last_name = extract_name(text)
            goals = extract_communication_goals(text, f"{first_name} {last_name}")
            max_goals = max(max_goals, len(goals))

            for i, g in enumerate(goals):
                max_subgoals_per_goal[i] = max(max_subgoals_per_goal.get(i, 0), len(g['subgoals']))

            all_data.append({
                'first_name': first_name,
                'last_name': last_name,
                'id': student_id,
                'goals': goals
            })

    columns = ['First Name', 'Last Name', 'Student ID']
    for g_idx in range(max_goals):
        columns.append(f'Goal {g_idx+1}')
        for s_idx in range(max_subgoals_per_goal.get(g_idx, 0)):
            columns.append(f'Benchmark {g_idx+1}.{s_idx+1}')

    rows = []
    for student in all_data:
        row = {
            'First Name': student['first_name'],
            'Last Name': student['last_name'],
            'Student ID': student['id']
        }
        for g_idx in range(max_goals):
            if g_idx < len(student['goals']):
                goal = student['goals'][g_idx]
                row[f'Goal {g_idx+1}'] = goal['goal']
                for s_idx in range(max_subgoals_per_goal[g_idx]):
                    row[f'Benchmark {g_idx+1}.{s_idx+1}'] = goal['subgoals'][s_idx] if s_idx < len(goal['subgoals']) else ''
            else:
                row[f'Goal {g_idx+1}'] = ''
                for s_idx in range(max_subgoals_per_goal[g_idx]):
                    row[f'Benchmark {g_idx+1}.{s_idx+1}'] = ''
        rows.append(row)

    df = pd.DataFrame(rows, columns=columns)
    df.to_excel(output_excel_path, index=False)
    return output_excel_path

# ------------------------
# DOCX FOLDER GENERATION
# ------------------------

def extract_id(text):
    match = re.search(r"Student ID:\s*(\d+)", text)
    return match.group(1) if match else None

def clean_action(text):
    text = text.strip().replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[.]+", "", text)
    match = re.search(r"will\s+(.*?)(?=\s*(with|given|in)\b|\d{1,3}%|$)", text, re.IGNORECASE)
    return match.group(1).strip().capitalize() if match else None

def create_goal_doc(goals_data, folder_path):
    doc = Document()
    for i, item in enumerate(goals_data, 1):
        doc.add_heading(f"Goal {i}", level=1)
        doc.add_paragraph(item['goal'])
        if item['subgoals']:
            doc.add_heading("Benchmarks", level=2)
            for b in item['subgoals']:
                doc.add_paragraph(b)
    doc.save(os.path.join(folder_path, "goals.docx"))

def create_note_doc(first_name, goals_data, folder_path):
    doc = Document()
    doc.add_heading("Note", level=1)
    doc.add_paragraph("Student was pulled-out and treated in therapy room setting and actively participated in an activity centered around a non-fiction grade level passage about the [ topic of passage ]")
    doc.add_paragraph("Utilizing the passage:")
    for item in goals_data:
        for i, b in enumerate(item['subgoals']):
            action = clean_action(b)
            if action:
                doc.add_paragraph(f"{first_name} achieved [ subgoal {i+1} percentage accuracy ] accuracy in being able to {action}.")
    doc.add_paragraph(f"{first_name} will continue current plan and using context clues to determine meaning.")
    doc.save(os.path.join(folder_path, "note.docx"))

def generate_docx_files(folder):
    excel_path = os.path.join(folder, "iep_goals_summary.xlsx")
    if not os.path.exists(excel_path):
        messagebox.showerror("Error", "iep_goals_summary.xlsx not found.")
        return

    df = pd.read_excel(excel_path)
    name_lookup = {
        str(row['Student ID']): (str(row['First Name']), str(row['Last Name']))
        for _, row in df.iterrows()
    }

    for filename in os.listdir(folder):
        if not filename.lower().endswith(".pdf"):
            continue

        pdf_path = os.path.join(folder, filename)
        reader = PdfReader(pdf_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)

        student_id = extract_id(text)
        if not student_id or student_id not in name_lookup:
            first_name, last_name = "Unknown", "Student"
        else:
            first_name, last_name = name_lookup[student_id]

        safe_name = f"{first_name}_{last_name}".replace(" ", "_")
        folder_name = f"{safe_name}_{student_id}"
        student_folder = os.path.join(folder, folder_name)
        os.makedirs(student_folder, exist_ok=True)

        goals_data = extract_communication_goals(text, f"{first_name} {last_name}")
        create_goal_doc(goals_data, student_folder)
        create_note_doc(first_name, goals_data, student_folder)

        shutil.move(pdf_path, os.path.join(student_folder, filename))

# ------------------------
# GUI
# ------------------------

def automate_day_one():
    folder = folder_var.get()
    if not folder:
        messagebox.showwarning("No Folder Selected", "Please select a folder to continue.")
        return

    try:
        output_excel = os.path.join(folder, "iep_goals_summary.xlsx")
        analyze_pdfs_and_generate_excel(folder, output_excel)
        generate_docx_files(folder)
        messagebox.showinfo("Success", "Excel and DOCX extraction completed.")
    except Exception as e:
        messagebox.showerror("Error", str(e))

root = tk.Tk()
root.title("IEP Day One Automation")
root.geometry("600x250")

folder_var = tk.StringVar()

tk.Label(root, text="Folder with IEP PDFs").pack(pady=5)
tk.Entry(root, textvariable=folder_var, width=60).pack(pady=5)
tk.Button(root, text="Browse", command=lambda: folder_var.set(filedialog.askdirectory())).pack(pady=5)
tk.Button(root, text="Automate Day One", command=automate_day_one, height=2).pack(pady=10)

root.mainloop()
