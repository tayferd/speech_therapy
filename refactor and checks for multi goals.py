import os
import re
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox
from PyPDF2 import PdfReader
from docx import Document

def extract_name(text):
    match = re.search(r"Student:\s+([A-Z ,'-]+)", text)
    return match.group(1).title().replace(",", "").replace("  ", " ") if match else None

def extract_id(text):
    match = re.search(r"Student ID:\s*(\d+)", text)
    return match.group(1) if match else None

def extract_first_name(name):
    if not name or name.lower().startswith("unknown"):
        return "Student"
    return name.split()[0].strip().title()

def clean_action(text):
    text = text.strip().replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[.]+", "", text)
    match = re.search(r"will\s+(.*?)(?=\s*(with|given|in)\b|\d{1,3}%|$)", text, re.IGNORECASE)
    return match.group(1).strip().capitalize() if match else None


def extract_communication_goals(text, name):
    import re

    comm_blocks = re.findall(
        r"Domain\(s\)/TSAA\(s\):\s*Communication(.*?)(?=\nDomain\(s\)/TSAA\(s\):|\nAssessments|\nAccommodations|Page \d|\Z)",
        text, re.DOTALL | re.IGNORECASE
    )

    results = []

    for section in comm_blocks:
        goal_blocks = re.split(r"\bGoal:\s*", section, flags=re.IGNORECASE)

        for block in goal_blocks[1:]:
            goal_text = re.split(
                r"(?:\nShort-term Objectives|Assessment Procedures|Progress Reported|\n[A-Z][a-z]+:|\Z)",
                block, maxsplit=1, flags=re.IGNORECASE
            )[0].strip().replace('\n', ' ')

            # FIRST: try to extract formally marked benchmarks
            benchmark_blocks = re.findall(
                r"Short-term Objectives or Benchmarks:(.*?)(?=\n(?:Short-term Objectives or Benchmarks:|Goal:|Domain\(s\)|Assessments|Accommodations|Page \d|\Z))",
                block, re.DOTALL | re.IGNORECASE
            )

            subgoals = []
            for bblock in benchmark_blocks:
                lines = bblock.strip().splitlines()
                for line in lines:
                    line = line.strip()
                    if (
                        len(line.split()) > 5 and
                        " will " in line.lower() and
                        not re.search(r"(student be|participate|assessment|instruction)", line, re.IGNORECASE)
                    ):
                        subgoals.append(line)

            # IF NO FORMAL BENCHMARKS, fallback to find will-statements as subgoals
            if not subgoals:
                lines = block.strip().splitlines()
                for line in lines:
                    line = line.strip()
                    if (
                        len(line.split()) > 5 and
                        " will " in line.lower() and
                        not re.search(r"(student be|participate|assessment|instruction)", line, re.IGNORECASE)
                    ):
                        subgoals.append(line)

            if goal_text:
                results.append({
                    "goal": goal_text,
                    "subgoals": subgoals
                })

    return results

def create_goal_doc(goals_data, folder_path):
    doc = Document()
    for i, item in enumerate(goals_data, 1):
        doc.add_heading(f"Goal {i}", level=1)
        doc.add_paragraph(item['goal'])
        if item['subgoals']:
            doc.add_heading("Benchmarks", level=2)
            for b in item['subgoals']:
                doc.add_paragraph(b)
    filename = os.path.join(folder_path, f"goals.docx")
    doc.save(filename)

def create_note_doc(first_name, goals_data, folder_path):
    doc = Document()
    doc.add_heading("Note", level=1)
    doc.add_paragraph("Student was pulled-out and treated in therapy room setting and actively participated in an activity centered around a non-fiction grade level passage about the [ topic of passage ]")
    doc.add_paragraph("Utilizing the passage:")

    for item in goals_data:
        for i, b in enumerate(item['subgoals']):
            action = clean_action(b)
            if action:
                doc.add_paragraph(f"{first_name} achieved [ subgoal {i+1} percentage accuracy ] accuracy in being able to {action}.")
    doc.add_paragraph(f"{first_name} will continue current plan and using context clues to determine meaning.")
    filename = os.path.join(folder_path, f"note.docx")
    doc.save(filename)

def process_pdfs(folder):
    for file_name in os.listdir(folder):
        if not file_name.lower().endswith(".pdf"):
            continue

        pdf_path = os.path.join(folder, file_name)
        reader = PdfReader(pdf_path)
        full_text = "\n".join(page.extract_text() or "" for page in reader.pages)

        name = extract_name(full_text)
        student_id = extract_id(full_text)
        goals_data = extract_communication_goals(full_text, name)

        if not name and goals_data:
            name_guess = re.search(r"([A-Z][a-z]+) will", goals_data[0]['goal'])
            name = f"Unknown {name_guess.group(1)}" if name_guess else "Unknown"

        safe_name = re.sub(r'[\\/*?:"<>|\n]', '', name.replace(",", "").replace(" ", "_"))
        folder_name = f"{safe_name}_{student_id or 'NoID'}"
        student_folder = os.path.join(folder, folder_name)
        os.makedirs(student_folder, exist_ok=True)

        first_name = extract_first_name(name)

        # ✅ New correct calls
        create_goal_doc(goals_data, student_folder)
        create_note_doc(first_name, goals_data, student_folder)

        shutil.move(pdf_path, os.path.join(student_folder, file_name))

def select_folder():
    folder = filedialog.askdirectory()
    folder_path_var.set(folder)

def run_extraction():
    folder = folder_path_var.get()
    if folder:
        process_pdfs(folder)
        messagebox.showinfo("Done", "PDFs processed successfully!")
    else:
        messagebox.showerror("Error", "Please select a folder first.")

# GUI
root = tk.Tk()
root.title("IEP Goal Extractor")
root.geometry("600x400")

folder_path_var = tk.StringVar()

tk.Label(root, text="Step 1: Select Folder with IEP PDFs").pack(pady=5)
tk.Button(root, text="Select Folder", command=select_folder).pack(pady=5)
tk.Label(root, textvariable=folder_path_var).pack(pady=5)
tk.Label(root, text="Step 2: Run Extraction").pack(pady=5)
tk.Button(root, text="Run Extraction", command=run_extraction).pack(pady=5)

root.mainloop()
